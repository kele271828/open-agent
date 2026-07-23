import requests
import psutil
import platform
import os
import time
import json
from pathlib import Path
from datetime import datetime, timedelta
import alibabacloud_oss_v2 as oss
from alibabacloud_oss_v2.credentials import StaticCredentialsProvider
import pyautogui
from ddgs import DDGS
import itertools
from docx import Document
import fitz  # PyMuPDF
import cv2
import pyperclip
from typing import Dict, Optional, Any, List  
import xml.etree.ElementTree as ET
from urllib.parse import urlparse
import subprocess
import asyncio
import winsdk.windows.devices.geolocation as wdg
import math
import uuid

from memory import SearchMemory, GetRecentMemory
# from settings import *
from config import config

# 开启安全防线：当鼠标移到屏幕四个角时，自动引发 pyautogui.FailSafeException 中断程序
pyautogui.FAILSAFE = True


def search_deep_memory(search_key = "", n_results = 3):
    if not search_key:
        return str(GetRecentMemory(limit=n_results, type="history")) # TODO: 处理返回值的格式
    # 搜索相关历史对话内容
    history_context = ""
    history = SearchMemory(search_key, n_results=n_results, type="history", date=None, distance_threshold=0.7)
    if history['metadatas']:
        for metadatas in history['metadatas']:
            if metadatas['date'] and metadatas['original_context']:
                history_context += f"{metadatas['date']} {metadatas['original_context']}\n"
    
    if history_context:
        return f"以下是相关历史对话内容，可能会有帮助。\n{history_context}\n"
    else:
        return "暂未发现相关历史对话"

            
def get_system_status() -> dict:
    """
    获取当前系统的核心资源状态（包含所有磁盘），兼容 Windows 和 Linux。
    返回结构化的字典供 AI 智能体解析。
    """
    status = {
        "os": platform.system(),
        "os_release": platform.release(),
        "cpu": {},
        "memory": {},
        "disks": {}  # 存储所有物理磁盘的信息
    }

    psutil.cpu_percent(interval=None)
    time.sleep(0.1)

    # 1. CPU 状态
    # 使用 interval=None 避免阻塞主线程。注意：首次调用会返回 0.0，建议在程序初始化时先调用一次。
    status["cpu"]["usage_percent"] = psutil.cpu_percent(interval=None)
    status["cpu"]["logical_cores"] = psutil.cpu_count(logical=True)
    status["cpu"]["physical_cores"] = psutil.cpu_count(logical=False)
    
    cpu_freq = psutil.cpu_freq()
    if cpu_freq:
        status["cpu"]["current_freq_mhz"] = round(cpu_freq.current, 2)

    # 2. 内存状态
    mem = psutil.virtual_memory()
    status["memory"]["total_gb"] = round(mem.total / (1024 ** 3), 2)
    status["memory"]["used_gb"] = round(mem.used / (1024 ** 3), 2)
    status["memory"]["usage_percent"] = mem.percent

    # 3. 遍历所有磁盘分区
    # all=False 参数非常重要，它会过滤掉 Linux 系统下的 /proc, /sys 等虚拟挂载点
    partitions = psutil.disk_partitions(all=False)
    
    for partition in partitions:
        try:
            # 尝试获取挂载点的空间使用情况
            usage = psutil.disk_usage(partition.mountpoint)
            
            # 以挂载点（如 'C:\\', 'D:\\' 或 '/'）作为字典的键
            status["disks"][partition.mountpoint] = {
                "device": partition.device,          # 物理设备路径，如 /dev/sda1
                "fstype": partition.fstype,          # 文件系统类型，如 NTFS, ext4
                "total_gb": round(usage.total / (1024 ** 3), 2),
                "used_gb": round(usage.used / (1024 ** 3), 2),
                "free_gb": round(usage.free / (1024 ** 3), 2),
                "usage_percent": usage.percent
            }
        except PermissionError:
            # 忽略没有权限访问的磁盘（如 Windows 下未插入介质的光驱或读卡器）
            pass
        except OSError:
            # 忽略由于设备未准备好引起的其他系统错误
            pass

    return json.dumps(status, ensure_ascii=False)




async def get_win11_location():
    # 1. 请求 Windows 系统的定位权限
    access_status = await wdg.Geolocator.request_access_async()
    
    if access_status == wdg.GeolocationAccessStatus.ALLOWED:
        print("已获得系统定位权限，正在获取位置...")
        
        # 2. 初始化定位器
        geolocator = wdg.Geolocator()
        
        try:
            # 3. 获取当前地理位置信息
            # 可以传入超时时间，这里默认等待硬件返回
            pos = await geolocator.get_geoposition_async()
            
            # 4. 提取经纬度和精度
            latitude = pos.coordinate.latitude
            longitude = pos.coordinate.longitude
            accuracy = pos.coordinate.accuracy
            # print("海拔：",pos.coordinate.altitude)
            return latitude, longitude, accuracy
            
        except Exception as e:
            return f"获取位置时发生错误: {e}"
            
    elif access_status == wdg.GeolocationAccessStatus.DENIED:
        return "错误：定位权限被拒绝。请在 Windows 设置中允许桌面应用访问位置。"
    else:
        return "错误：无法明确的权限状态。"

def get_location():
    location = asyncio.run(get_win11_location())
    if isinstance(location, tuple):
        return f"纬度：{location[0]:.6f}，经度：{location[1]:.6f}，定位精度：{location[2]:.2f}米"
    else:
        return location

def haversine_distance(lat1, lon1, lat2, lon2):
    """
    计算两个经纬度点之间的球面距离（单位：米）
    """
    R = 6371000.0  # 地球平均半径，单位为米

    # 将十进制的度数转换为弧度
    phi1 = math.radians(lat1)
    phi2 = math.radians(lat2)
    delta_phi = math.radians(lat2 - lat1)
    delta_lambda = math.radians(lon2 - lon1)

    # Haversine 核心公式
    a = math.sin(delta_phi / 2.0) ** 2 + \
        math.cos(phi1) * math.cos(phi2) * \
        math.sin(delta_lambda / 2.0) ** 2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))

    distance = R * c
    return distance

def get_semantic_location(current_lat, current_lon, threshold_meters=80):
    """
    将当前坐标与预设地点进行匹配
    threshold_meters: 匹配的判定半径（误差容忍度，默认80米）
    """
    closest_place = "未知地点 (不在预设范围内)"
    min_distance = float('inf')

    # 遍历所有预设地点，寻找距离最近的一个
    for place_name, (place_lat, place_lon) in config.MY_LOCATIONS.items():
        dist = haversine_distance(current_lat, current_lon, place_lat, place_lon)
        print(f"距离 {place_name} 的距离: {dist:.2f} 米")
        if dist < min_distance:
            min_distance = dist
            # 如果最短距离在你设定的容忍半径内，则判定匹配成功
            if dist <= threshold_meters:
                closest_place = place_name

    return closest_place, min_distance

def check_location():
    location = asyncio.run(get_win11_location())
    if isinstance(location, tuple):
        my_lat, my_lon, accuracy = location
    else:
        return location
    
    place, distance = get_semantic_location(my_lat, my_lon, threshold_meters=max(accuracy, 100))
    
    print(f"当前经纬度: {my_lat}, {my_lon}, 精度：{accuracy:.2f}米")
    if place != "未知地点 (不在预设范围内)":
        return f"可能在：{place} (距离预设中心点约 {distance:.1f} 米)"
    else:
        return f"未能精确匹配。距离最近的已知地点还有 {distance:.1f} 米。"


def capture_single_image(save_path="tmp/image/vision_capture.jpg"):
    """
    调用本地摄像头拍摄单张照片，保存并返回图像数据。
    
    参数:
        save_path (str): 照片保存的文件路径和名称。
        
    返回:
        frame (numpy.ndarray or None): 拍摄到的图像帧数据，如果失败则返回 None。
    """
    # 0 代表系统默认的第一个摄像头
    cap = cv2.VideoCapture(0)

    if not cap.isOpened():
        return "错误：无法打开摄像头，请检查占用或系统隐私设置。"
        return None

    print("正在启动摄像头并调整曝光，请稍候...")

    # ==========================================
    # 核心技巧：摄像头预热
    # 连续读取并丢弃前 10 帧，给硬件留出自动调光的时间
    # ==========================================
    for _ in range(10):
        cap.read()
        time.sleep(0.05) # 短暂休眠，等待硬件响应

    # 预热完毕，正式拍摄我们要的那一张
    ret, frame = cap.read()

    if ret:
        # 将图片保存到本地磁盘
        cv2.imwrite(save_path, frame)
    else:
        return "错误：成功打开了摄像头，但无法获取画面。"
        frame = None

    # 无论成功与否，务必释放摄像头资源，以免造成硬件持续开启（指示灯长亮）
    cap.release()

    return [{"type":"text","text":f"拍摄成功，照片已保存至: {save_path}"},
             {"type":"image_url","image_url":upload_file_and_get_url(save_path)}]



def get_current_time():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S %A")




def get_weather(location=""):
    """
    获取指定城市的天气预报，并浓缩为最适合 LLM 读取的精简 JSON 格式。
    参数 location: 城市名（英文拼音或中文均可，留空则自动定位）
    """
    try:
        # 使用 j1 获取 JSON 数据
        url = f"https://wttr.in/{location}?format=j1"
        # 加上 timeout 防止工具卡死导致 LLM 响应超时
        response = requests.get(url, timeout=10)
        
        if response.status_code == 200:
            raw_data = response.json()
            
            # 1. 初始化精简版的数据字典
            llm_data = {
                "location": location if location else "自动定位当前位置",
                "current": {},
                "forecast": []
            }
            
            # 2. 提取当前天气核心数据
            current = raw_data.get('current_condition', [{}])[0]
            llm_data["current"] = {
                "temp_C": current.get('temp_C'),
                "humidity": current.get('humidity'),
                "condition": current.get('weatherDesc', [{}])[0].get('value')
            }
            
            # 3. 提取未来三天预报（每天只取最高温、最低温和大致天气）
            for day in raw_data.get('weather', []):
                # wttr.in 每天有 8 个时段的数据，我们取中午 12:00 (索引4) 的天气描述作为当天概览
                midday_condition = day.get('hourly', [{}])[4].get('weatherDesc', [{}])[0].get('value')
                
                llm_data["forecast"].append({
                    "date": day.get('date'),
                    "max_C": day.get('maxtempC'),
                    "min_C": day.get('mintempC'),
                    "condition": midday_condition
                })
                
            # 4. 转换为最紧凑的 JSON 字符串
            # separators=(',', ':') 可以去掉所有多余空格，进一步节省 Token
            # ensure_ascii=False 确保中文正常显示，不被转义成极占 Token 的 Unicode 编码
            return json.dumps(llm_data, ensure_ascii=False, separators=(',', ':'))
            
        else:
            return json.dumps({"error": f"API请求失败，状态码: {response.status_code}"}, ensure_ascii=False)
            
    except Exception as e:
        # 捕获所有异常并以 JSON 格式返回，防止大模型工具链崩溃
        return json.dumps({"error": f"发生内部错误: {str(e)}"}, ensure_ascii=False)


def web_search_for_llm(query, max_results=5):
    """
    为 AI 助手设计的联网搜索函数
    返回格式化的字符串，方便直接喂给 LLM
    """
    search_data = []
    try:
        # 使用 DDGS 上下文管理器，自动管理连接
        with DDGS() as ddgs:
            # 获取文本搜索结果
            # region="wt-wt" 为全球，"cn-zh" 为中国地区（如果需要中文优化）
            results = ddgs.text(query, region="wt-wt", safesearch="moderate", timelimit="y")
            
            # 仅取前 max_results 条结果
            for i, r in enumerate(itertools.islice(results, max_results)):
                # 拼接成 LLM 易读的格式
                entry = f"[{i+1}] 来源: {r['href']}\n标题: {r['title']}\n摘要: {r['body']}\n"
                search_data.append(entry)
                
        if not search_data:
            return "未找到相关联网信息。"
            
        return "\n".join(search_data)

    except Exception as e:
        return f"联网搜索时发生错误: {str(e)}"



def query_academic_papers(query: str, author: str = None, year: str = None, limit: int = 8) -> str:
    """
    使用无需注册的 ArXiv API 查询学术论文
    """
    url = "http://export.arxiv.org/api/query"
    
    # 构造 ArXiv 特有的查询字符串
    # 默认在所有字段中搜索 (all:query)
    search_query = f"all:{query}"
    # search_query = f'ti:"{query}"'
    
    # 如果指定了作者，可以拼接到搜索条件中
    if author:
        search_query += f" AND au:{author}"
        
    params = {
        "search_query": search_query,
        "start": 0,
        "max_results": limit,
        "sortBy": "relevance", # 按相关度排序
        "sortOrder": "descending"
    }
    
    try:
        response = requests.get(url, params=params, timeout=10)
        response.raise_for_status()
        
        # 解析返回的 XML 数据
        root = ET.fromstring(response.text)
        
        # ArXiv 使用了 Atom 命名空间
        ns = {'atom': 'http://www.w3.org/2005/Atom'}
        entries = root.findall('atom:entry', ns)
        
        if not entries:
            return json.dumps({"status": "success", "message": "未找到相关论文", "results": []})
            
        results = []
        for entry in entries:
            # 提取标题并清理换行符
            title = entry.find('atom:title', ns).text.replace('\n', ' ').strip()
            
            # 提取摘要
            summary = entry.find('atom:summary', ns).text.replace('\n', ' ').strip()
            
            # 提取发表年份 (ArXiv 返回的是如 2017-06-12T17:33:52Z 的格式)
            published = entry.find('atom:published', ns).text
            pub_year = published[:4] if published else "Unknown"
            
            # 提取链接
            paper_url = entry.find('atom:id', ns).text
            
            # 提取作者列表
            authors = [author.find('atom:name', ns).text for author in entry.findall('atom:author', ns)]
            
            # 如果用户指定了年份，可以在这里做简单的本地过滤
            if year and year != pub_year:
                continue
                
            results.append({
                "title": title,
                "authors": ", ".join(authors),
                "year": pub_year,
                "abstract": summary[:500] + "...", # 截断摘要
                "url": paper_url,
                "venue": "ArXiv"
            })
            
        # 如果本地按年份过滤后结果为空
        if not results:
             return json.dumps({"status": "success", "message": f"未找到 {year} 年的匹配论文", "results": []})
             
        # 返回前 limit 个结果（以防本地过滤后数量变少）
        return json.dumps({"status": "success", "results": results[:limit]}, ensure_ascii=False)
        
    except requests.exceptions.RequestException as e:
        return json.dumps({"status": "error", "message": f"网络请求失败: {str(e)}"})
    except ET.ParseError:
        return json.dumps({"status": "error", "message": "解析 API 返回数据失败"})



def read_academic_paper(arxiv_url: str, download_dir: str = "tmp/download_papers") -> str:
    """
    根据 ArXiv 链接下载论文 PDF 到指定目录，并提取全文文本
    
    :param arxiv_url: 论文的 ArXiv 链接 (例如: http://arxiv.org/abs/2105.02723v1)
    :param download_dir: 本地保存目录，默认为当前路径下的 papers 文件夹
    """
    try:
        # 1. 转换 URL 为 PDF 直链
        if "arxiv.org/abs/" in arxiv_url:
            pdf_url = arxiv_url.replace("/abs/", "/pdf/") + ".pdf"
        elif "arxiv.org/pdf/" in arxiv_url:
            pdf_url = arxiv_url if arxiv_url.endswith(".pdf") else arxiv_url + ".pdf"
        else:
            return json.dumps({"status": "error", "message": "目前仅支持解析 ArXiv 链接"})

        # 2. 提取文件名并准备本地路径
        parsed_url = urlparse(pdf_url)
        filename = os.path.basename(parsed_url.path)
        if not filename:
            filename = "downloaded_paper.pdf"
            
        # 确保下载目录存在，如果不存在则自动创建
        os.makedirs(download_dir, exist_ok=True)
        file_path = os.path.join(download_dir, filename)

        # 3. 下载 PDF 文件
        headers = {'User-Agent': 'Mozilla/5.0'} 
        response = requests.get(pdf_url, headers=headers, timeout=15)
        response.raise_for_status()

        # 4. 将 PDF 数据写入本地文件
        with open(file_path, "wb") as f:
            f.write(response.content)

        # 5. 读取本地文件并解析文本
        # 使用 fitz 打开刚刚保存的本地文件
        doc = fitz.open(file_path)
        
        full_text = ""
        for page in doc:
            full_text += page.get_text()
            
        doc.close() # 解析完毕后关闭文件句柄

        # 6. 清理文本并限制长度
        full_text = full_text.replace("-\n", "").replace("\n", " ")
        truncated_text = full_text[:20000] # 截断前 20000 字符

        # 7. 在返回结果中增加 file_path，让大模型知道文件存在哪里
        return json.dumps({
            "status": "success", 
            "message": f"论文已成功下载至 {file_path} 并完成解析", 
            "file_path": file_path,
            "content": truncated_text
        }, ensure_ascii=False)

    except requests.exceptions.RequestException as e:
        return json.dumps({"status": "error", "message": f"PDF 下载失败: {str(e)}"})
    except Exception as e:
        return json.dumps({"status": "error", "message": f"PDF 处理失败: {str(e)}"})




def take_screenshot(save_path="tmp/image/screenshot.png", region=None):
    """
    跨平台截屏函数。优先使用 pyautogui。
    
    :param save_path: 截图保存的路径和文件名 (如 'test.png')。
    :param region: 截图区域 (左, 上, 宽, 高) 的元组。如果为 None 则截取全屏。
    :return: 截图的图片对象 (PIL.Image)，如果失败则返回 None。
    """
    try:
        
        # pyautogui.screenshot() 接受 region 格式为 (left, top, width, height)
        if region:
            img = pyautogui.screenshot(region=region)
        else:
            img = pyautogui.screenshot()
            
        img.save(save_path)
        
        return [{"type": "text", "text": f"截图成功，保存至 {os.path.abspath(save_path)}，分辨率：{img.size}"},
                {"type": "image_url", "image_url": upload_file_and_get_url(save_path,filename=f"{uuid.uuid4().hex}.png")}]
        
    except Exception as e:
        print(f"❌ 截屏过程中发生错误: {e}")
        return None



def control_mouse(action: str, x: int = None, y: int = None, button: str = 'left', clicks: int = 1, amount: int = 0, duration: float = 0.25) -> dict:
    """
    AI助手专用的单节点鼠标控制接口。
    
    :param action: 执行的动作，支持: 'position' (获取坐标), 'size' (获取屏幕大小), 'move' (移动), 'click' (点击), 'drag' (拖拽), 'scroll' (滚动)
    :param x: 目标X坐标 (move, click, drag 需要)
    :param y: 目标Y坐标 (move, click, drag 需要)
    :param button: 鼠标按键 ('left', 'right', 'middle')
    :param clicks: 点击次数
    :param amount: 滚动幅度 (正数向上，负数向下，scroll 需要)
    :param duration: 动作持续时间(秒)，使移动平滑避免触发风控
    :return: 包含操作状态和结果的字典
    """
    
    try:
        if action == 'position':
            pos_x, pos_y = pyautogui.position()
            return f"成功获取当前鼠标位置: x: {pos_x}, y: {pos_y}"
            
        elif action == 'size':
            width, height = pyautogui.size()
            return f"成功获取屏幕尺寸: width: {width}, height: {height}"
            
        elif action == 'move':
            if x is None or y is None:
                return f"移动操作缺少 x 或 y 坐标"
            pyautogui.moveTo(x, y, duration=duration)
            return f"成功移动到 ({x}, {y})"
            
        elif action == 'click':
            # 如果提供了坐标，先移动过去再点；否则在原地点击
            if x is not None and y is not None:
                pyautogui.moveTo(x, y, duration=duration)
            pyautogui.click(button=button, clicks=clicks)
            return f"执行了 {clicks} 次 {button} 键点击"
            
        elif action == 'drag':
            if x is None or y is None:
                return f"拖拽操作缺少 x 或 y 坐标"
            pyautogui.dragTo(x, y, duration=duration, button=button)
            return f"按住 {button} 键拖拽到了 ({x}, {y})"
            
        elif action == 'scroll':
            if amount == 0:
                return f"滚动操作缺少 amount 滚动量参数"
            pyautogui.scroll(amount)
            return f"已滚动 {amount} 个单位"
            
        else:
            return f"未知的指令: {action}"
            
    except pyautogui.FailSafeException:
        return f"触发防故障：鼠标移动到了屏幕角落"
    except Exception as e:
        return f"操作异常: {str(e)}"
    


def control_keyboard(action: str, text: str = "", keys: list = None, interval: float = 0.05) -> str:
    """
    控制键盘的统一接口，适用于 AI Agent 调用。
    
    参数:
        action (str): 执行的动作类型。可选值:
                      - 'write': 输入纯英文/数字文本
                      - 'write_cn': 输入包含中文的文本（通过剪贴板粘贴）
                      - 'press': 按下单个或多个按键（如 Enter, Tab）
                      - 'hotkey': 执行快捷键组合（如 Ctrl+C）
        text (str): 需要输入的文本内容。
        keys (list): 需要按下的键名列表，例如 ['enter'], ['ctrl', 'c']。
        interval (float): 连续输入时的按键间隔时间（秒）。
        
    返回:
        str: 执行结果的状态描述。
    """
    try:
        # if action == "write":
        #     if not text:
        #         return "Error: 'text' 参数不能为空。"
        #     pyautogui.write(text, interval=interval)
        #     return f"Success: 已输入文本 '{text}'"

        if action == "write":
            if not text:
                return "Error: 'text' 参数不能为空。"
            # 利用剪贴板实现中文输入
            # 保存当前剪贴板内容，防止覆盖
            original_clipboard = pyperclip.paste()
            pyperclip.copy(text)
            time.sleep(0.1) # 稍微等待剪贴板写入
            # 兼容 Mac 和 Windows 的粘贴快捷键
            import platform
            if platform.system() == 'Darwin':
                pyautogui.hotkey('command', 'v')
            else:
                pyautogui.hotkey('ctrl', 'v')
            # 恢复原始剪贴板内容
            pyperclip.copy(original_clipboard)
            return f"Success: 已通过粘贴输入中文文本 '{text}'"

        elif action == "press":
            if not keys:
                return "Error: 'keys' 参数不能为空。"
            for key in keys:
                pyautogui.press(key)
                time.sleep(interval)
            return f"Success: 已依次按下按键 {keys}"

        elif action == "hotkey":
            if not keys or len(keys) < 2:
                return "Error: 'hotkey' 动作至少需要两个按键。"
            # 解包列表传入 hotkey
            pyautogui.hotkey(*keys)
            return f"Success: 已触发快捷键 {'+'.join(keys)}"

        else:
            return f"Error: 未知的动作类型 '{action}'。"

    except pyautogui.FailSafeException:
        return "Error: 触发了 PyAutoGUI 安全防线（鼠标移至屏幕角落），操作已终止。"
    except Exception as e:
        return f"Error: 键盘控制执行失败: {str(e)}"
    


def manage_clipboard(action: str, text: str = "") -> str:
    """
    管理系统剪贴板的统一接口。
    
    参数:
        action (str): 执行的动作类型。可选值:
                      - 'read': 读取当前剪贴板的文本内容
                      - 'write': 将指定文本写入剪贴板
        text (str): 当 action 为 'write' 时，需要写入剪贴板的内容。
        
    返回:
        str: 如果是 'read'，返回剪贴板中的文本（若为空或非纯文本则返回提示语）。
             如果是 'write'，返回操作状态信息。
    """
    try:
        if action == "read":
            # 获取剪贴板内容
            content = pyperclip.paste()
            # 过滤掉空字符串或全是空格的情况
            if not content or not content.strip():
                return "Clipboard is empty or contains non-text data (e.g., an image or file)."
            return content

        elif action == "write":
            if not text:
                return "Error: 'text' parameter cannot be empty for write action."
            # 将文本写入剪贴板
            pyperclip.copy(text)
            return "Success: 文本已成功写入剪贴板。"

        else:
            return f"Error: 未知的动作类型 '{action}'。请使用 'read' 或 'write'。"

    except pyperclip.PyperclipException as e:
        # 处理跨平台剪贴板依赖缺失的问题
        return f"Error: 剪贴板环境异常。如果在 Linux 环境下，请确保已安装 xclip 或 xsel。详细报错: {str(e)}"
    except Exception as e:
        return f"Error: 剪贴板操作失败: {str(e)}"
    

def is_in_whitelist(target_path: str) -> bool:
    """检查路径是否命中了白名单规则（支持文件夹）"""
    for white_path in config.WHITELIST_PATHS:
        # 如果目标路径正好是白名单文件，或者位于白名单目录下（加 os.sep 确保是子目录匹配）
        if target_path == white_path or target_path.startswith(white_path + os.sep):
            return True
    return False

def is_in_blacklist(target_path: str) -> bool:
    """检查路径是否命中了黑名单规则"""
    for black_path in config.BLACKLIST_PATHS:
        # 如果目标路径正好是黑名单文件，或者位于黑名单目录下（加 os.sep 确保是子目录匹配）
        if target_path == black_path or target_path.startswith(black_path + os.sep):
            return True
    return False

def read_file_content(file_path: str):
    """
    读取文件内容（支持纯文本读取与图片上传生成URL）
    
    :param file_path: 文件路径
    :return: 文本内容字符串，或者包含图片URL的结构化列表
    """
    try:
        # 1. 规范化路径
        target_path = os.path.abspath(file_path)
        
        # 提取文件后缀，转换为小写备用
        _, ext = os.path.splitext(target_path)
        ext_lower = ext.lower()
        
        # 2. 【白名单校验】优先处理
        if is_in_whitelist(target_path):
            pass # 命中白名单，直接跳过后续的安全限制
            
        else:
            # 3. 【黑名单校验】
            if is_in_blacklist(target_path):
                return f"Error: 访问被拒绝，目标 {file_path} 位于黑名单中。"
            
            # 4. 【驱动器校验】限制在 D 盘内
            if not target_path.startswith(config.ALLOWED_DRIVE):
                return f"Error: 越权访问，仅允许访问 D 盘或白名单中的文件。"
                
            # 5. 【后缀名校验】同时允许纯文本后缀和图片后缀
            if ext_lower not in config.ALLOWED_EXTENSIONS and ext_lower not in config.ALLOWED_IMAGE_EXTENSIONS:
                return f"Error: 访问被拒绝，不支持读取 {ext} 类型的文件。"

        # 6. 检查文件物理状态
        if not os.path.exists(target_path):
            return f"Error: 文件 {file_path} 不存在。"
        if not os.path.isfile(target_path):
            return f"Error: {file_path} 是一个目录，而不是文件。"

        # 7. 根据文件类型分支处理：图片上传 or 纯文本读取
        # 即使是白名单里的文件，我们也需要通过后缀判断它是不是图片
        if ext_lower in config.ALLOWED_IMAGE_EXTENSIONS or ext_lower in {'.jpg', '.jpeg', '.png', '.gif', '.webp'}:
            # 构造 filename，例如拼接成 "test1.jpg"
            upload_filename = f"{uuid.uuid4().hex}{ext_lower}"
            
            # 调用上传函数获取 public_url
            public_url = upload_file_and_get_url(target_path, filename=upload_filename)
            
            # 返回符合大模型视觉API要求的列表格式
            return [
                {"type": "image_url", "image_url": {"url": public_url}},
                ]
        
        elif ext_lower in {'.docx'}:
            # 调用 Word 文档解析函数
            return read_word_with_index(target_path)
        
        elif ext_lower in {'.pdf'}:
            # 调用 PDF 解析函数
            return read_pdf(target_path)

        else:
            # 安全读取文本并返回内容
            with open(target_path, 'r', encoding='utf-8') as file:
                return file.read()
            
    except UnicodeDecodeError:
        return f"Error: 文件 {file_path} 编码可能不是 UTF-8，无法解析纯文本。"
    except Exception as e:
        return f"Error: 读取/上传文件 {file_path} 时发生未知错误: {e}"
    

def explore_directory(dir_path: str) -> str:
    """
    探路函数：让 AI 查看指定目录下的所有文件夹和文件列表。
    
    :param dir_path: 想要探索的目录路径
    :return: 格式化的目录内容文本
    """
    try:
        # 1. 规范化路径
        target_path = os.path.abspath(dir_path)
        
        # 2. 【黑名单校验】一票否决
        if is_in_blacklist(target_path):
            return f"Error: 视线被阻挡，目录 {dir_path} 位于黑名单中，禁止窥探。"
            
        # 3. 【白名单与驱动器校验】
        # 如果不在白名单文件夹内，就必须受限于 D 盘
        if not is_in_whitelist(target_path):
            if not target_path.startswith(config.ALLOWED_DRIVE):
                return f"Error: 越界探路，仅允许浏览 D 盘或白名单中的目录。"

        # 4. 检查物理状态
        if not os.path.exists(target_path):
            return f"Error: 目录 {dir_path} 不存在。"
        if not os.path.isdir(target_path):
            return f"Error: {dir_path} 是一个文件，不是目录。请使用读取文件功能。"

        # 5. 开始探路（获取列表）
        items = os.listdir(target_path)
        
        folders = []
        files = []
        
        for item in items:
            item_path = os.path.join(target_path, item)
            
            # 【安全过滤】如果在黑名单里，直接对 AI 隐身，不展示在列表中
            if is_in_blacklist(item_path):
                continue
                
            if os.path.isdir(item_path):
                # 用特殊符号或后缀标识文件夹，让 AI 知道这个路径可以继续 explore_directory
                folders.append(f"📁 [目录] {item}")
            else:
                # 标识文件，让 AI 知道这个路径需要用 read_file_content
                files.append(f"📄 [文件] {item}")
        
        # 6. 将结果格式化为 AI 容易理解的纯文本
        result = f"当前位置: {target_path}\n"
        result += "-" * 30 + "\n"
        result += "【包含的文件夹】:\n" + ("\n".join(folders) if folders else "  (无)") + "\n\n"
        result += "【包含的文件】:\n" + ("\n".join(files) if files else "  (无)")
        
        return result
        
    except PermissionError:
         return f"Error: 权限不足，无法查看 {dir_path} (可能是系统保护的隐藏文件夹)。"
    except Exception as e:
        return f"Error: 探路时发生未知错误: {e}"


def modify_file(file_path: str, action: str, content: str = "", old_text: str = "", overwrite: bool = False) -> str:
    """
    终极安全文件修改函数：支持写入、追加、局部替换。
    
    :param file_path: 目标文件路径
    :param action: 执行的动作，必须是 'write', 'append', 或 'replace'
    :param content: 要写入/追加的新内容，或替换时的新文本 (对应 new_text)
    :param old_text: 仅在 action='replace' 时需要，表示要被替换的旧文本
    :param overwrite: 仅在 action='write' 时有效，是否允许覆盖已存在的文件
    :return: 执行结果信息
    """
    try:
        # ================= 1. 统一的绝对安全防线 =================
        target_path = os.path.abspath(file_path)
        
        # 【黑名单校验】
        if is_in_blacklist(target_path):
            return f"Error: 操作被拒绝，目标 {file_path} 位于黑名单中。"
            
        # 【白名单与驱动器、后缀校验】
        if not is_in_whitelist(target_path):
            if not target_path.startswith(config.ALLOWED_DRIVE):
                return f"Error: 越界操作，仅允许在 D 盘或白名单目录中修改文件。"
                
            _, ext = os.path.splitext(target_path)
            if ext.lower() not in config.ALLOWED_WRITE_EXTENSIONS:
                return f"Error: 极其危险！禁止向 {ext} 格式的文件写入数据。"

        # 【内容大小限制】
        if len(content.encode('utf-8')) > config.MAX_WRITE_SIZE:
            return "Error: 写入内容过大，超出 5MB 限制。"

        # ================= 2. 自动创建父目录 =================
        target_dir = os.path.dirname(target_path)
        if not os.path.exists(target_dir):
            os.makedirs(target_dir, exist_ok=True)

        # ================= 3. 根据 Action 执行具体分支 =================
        
        # 分支 A: 全新写入 / 覆盖
        if action == "write":
            if os.path.exists(target_path):
                if not overwrite:
                    return f"Error: 文件 {file_path} 已存在。为了防止误删，默认禁止覆盖。如需覆盖，请设置 overwrite=True。"
                if not os.path.isfile(target_path):
                     return f"Error: 目标 {file_path} 是一个已存在的文件夹，无法写入。"
                     
            with open(target_path, 'w', encoding='utf-8') as file:
                file.write(content)
            return f"Success: 文件已成功写入至 {file_path}。"

        # 分支 B: 追加内容
        elif action == "append":
            if not os.path.exists(target_path):
                return f"Error: 文件 {file_path} 不存在，无法追加。请先使用 'write' 动作创建文件。"
                
            if not content.startswith('\n'):
                content = '\n' + content
                
            with open(target_path, 'a', encoding='utf-8') as file:
                file.write(content)
            return f"Success: 内容已成功追加至 {file_path}。"

        # 分支 C: 局部替换
        elif action == "replace":
            if not os.path.exists(target_path):
                return f"Error: 文件 {file_path} 不存在，无法替换内容。"
            if not old_text:
                return "Error: 执行替换操作时，必须提供 old_text 参数。"

            with open(target_path, 'r', encoding='utf-8') as file:
                original_content = file.read()

            if old_text not in original_content:
                return f"Error: 未在文件中找到完全匹配的旧文本，替换失败。请确保 old_text 逐字完全一致。"

            updated_content = original_content.replace(old_text, content)

            with open(target_path, 'w', encoding='utf-8') as file:
                file.write(updated_content)
            return f"Success: 文件 {file_path} 中的指定内容已成功更新。"

        # 未知指令兜底
        else:
            return f"Error: 未知的 action '{action}'。仅支持 'write', 'append', 'replace'。"

    except Exception as e:
        return f"Error: 修改文件时发生未知错误: {e}"
    

def read_word_with_index(file_path):
    doc = Document(file_path)
    structured_content = []
    
    for index, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue # 跳过纯空白段落，节省 Token
            
        # 提取样式，帮助 LLM 理解文档结构（是标题还是正文）
        style_name = para.style.name
        
        # 拼接成带索引的格式，例如: [3] (Heading 1) 这是一个大标题
        structured_content.append(f"[{index}] ({style_name}) {text}")
        
    return "\n".join(structured_content)


def read_pdf(file_path):
    text = ""
    # 打开 PDF 文件
    with fitz.open(file_path) as doc:
        # 遍历每一页提取文本
        for page in doc:
            text += page.get_text()
    return text


def modify_word_element(doc_path: str, action: str, key: int, text: str = None) -> str:
    """
    根据 LLM 提供的 action 和 key 修改 Word 文档。
    
    :param doc: python-docx 的 Document 对象
    :param action: 'MODIFY' (修改), 'ADD_AFTER' (在之后添加), 'DELETE' (删除)
    :param key: 目标段落的序号 (index)
    :param text: 提供的文本，默认为 None
    :return: bool，表示是否执行成功
    """
    # 获取当前文档的所有段落快照
    if os.path.exists(doc_path):
        doc = Document(doc_path)
    else:
        # 文件不存在时，如果不允许新建，可以直接 return 报错
        # 如果允许新建，则创建一个全新的空白文档对象
        doc = Document() 
        # 针对新建文档的特殊兜底：如果是新建的，通常直接把 text 加上去即可
        if text:
            doc.add_paragraph(text)
            doc.save(doc_path)
            return "文件不存在，已自动新建并写入了首段内容。"
        else:
            return "[执行失败] 文件不存在，且没有提供初始文本(text)。"
        
    paragraphs = doc.paragraphs
    
    # 1. 边界与合法性检查
    if key < 0 or key >= len(paragraphs):
        return f"[执行失败] 越界：找不到序号为 {key} 的段落。当前最大序号为 {len(paragraphs)-1}"
        
        
    target_para = paragraphs[key]
    action = action.upper() # 统一转大写，增加容错率
    
    try:
        # 2. 执行具体动作
        if action == 'MODIFY':
            if text is None:
                return f"[执行失败] MODIFY 操作缺少 text 参数 (key={key})"
                
            # 【核心改进】：获取原段落第一个 Run 的格式（如果存在的话）
            original_style = None
            if len(target_para.runs) > 0:
                first_run = target_para.runs[0]
                # 保存核心样式属性
                original_style = {
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                    'underline': first_run.underline,
                    'font_name': first_run.font.name if first_run.font else None,
                    'font_size': first_run.font.size if first_run.font else None,
                    'color': first_run.font.color.rgb if first_run.font and first_run.font.color else None
                }

            # 清空原段落（这保留了段落级别的整体样式，比如缩进、居中等）
            target_para.clear()
            
            # 添加新文本作为一个全新的 Run
            new_run = target_para.add_run(text)
            
            # 尝试恢复前面保存的样式
            if original_style:
                new_run.bold = original_style['bold']
                new_run.italic = original_style['italic']
                new_run.underline = original_style['underline']
                if original_style['font_name']:
                    new_run.font.name = original_style['font_name']
                if original_style['font_size']:
                    new_run.font.size = original_style['font_size']
                if original_style['color']:
                    new_run.font.color.rgb = original_style['color']
            
        elif action == 'DELETE':
            # 删除段落需要深入底层 XML 操作：找到父节点并移除当前节点
            p_element = target_para._element
            p_element.getparent().remove(p_element)
            target_para._element = None # 清理引用
            
        elif action == 'ADD_AFTER':
            if text is None:
                return f"[执行失败] ADD_AFTER 操作缺少 text 参数 (key={key})"
            # 先在文档末尾生成一个新段落
            new_p = doc.add_paragraph(text)
            # 通过操作底层 XML，将这个新生成的段落节点强行插队到目标段落之后
            target_para._element.addnext(new_p._element)
            
        else:
            return f"[执行失败] 未知的 action: {action}"
        
        doc.save(doc_path) # 保存修改
        return "修改成功"
        
    except Exception as e:
        return f"[执行异常] key={key}, action={action}, Error: {str(e)}"




def launch_app(app_name: str, external_mapping: Optional[Dict[str, str]] = config.APP_REGISTRY) -> str:
    """
    启动本地应用程序的函数式接口。
    优先检索外部映射表，如果未命中则直接尝试将其作为系统已知命令启动。
    
    参数:
        app_name: AI 输出的应用程序名称 (如 "微信", "edge", "chrome.exe")
        external_mapping: 外部映射表字典 {应用别名: 绝对路径或可执行文件名}
    """
    if platform.system() != "Windows":
        return "执行失败：os.startfile() 仅支持 Windows 系统。"
        
    # 1. 初始化空映射表以防未传入
    mapping = external_mapping or {}
    
    # 2. 预处理：将映射表的键名统一转为小写，实现无视大小写的模糊匹配
    normalized_mapping = {k.lower(): v for k, v in mapping.items()}
    search_key = app_name.strip().lower()

    # 3. 核心逻辑：优先查表，查不到则 fallback 到原始输入
    # 如果是绿色软件或微信，命中映射表返回绝对路径；如果是 edge，直接返回 app_name
    target_path = normalized_mapping.get(search_key, app_name)

    # 4. 执行操作与副作用隔离
    try:
        # os.startfile 会调用 Windows Shell，能解析绝对路径、App Paths 和 shell: 协议
        os.startfile(target_path)
        return f"已成功发送启动指令: {app_name}"
        
    except FileNotFoundError:
        return (f"启动失败：找不到目标 '{target_path}'。 "
                f"如果是第三方或绿色软件，请将其绝对路径添加到映射表中。")
    except Exception as e:
        return f"启动失败：执行时发生未知错误 ({str(e)})"



def safe_terminate_process(process_name):
    """
    基于白名单机制的进程关闭函数，供 AI 助手的 Tool/Function Calling 使用。
    
    :param process_name: AI 提取的进程名称 (如 "Code.exe", "QQ", "firefox")
    :return: 执行结果字符串，直接作为 Observation 返回给大模型
    """
    if not process_name:
        return "错误：未提供进程名称。"

    # 2. 清理输入参数 (跨平台兼容)
    # 无论传入 "Code.exe", "CODE", 还是 "code"，都统一处理为 "code"
    target_name = process_name.lower().replace(".exe", "")

    # 3. 拦截网关：白名单校验
    if target_name not in config.ALLOWED_PROCESSES:
        return f"权限拒绝：'{process_name}' 不在安全白名单中。系统仅允许关闭预设的安全应用。"

    system_type = platform.system()
    terminated_count = 0

    # 4. 执行关闭逻辑
    for proc in psutil.process_iter(['pid', 'name']):
        try:
            # 清理当前抓取到的进程名
            current_proc_name = proc.info['name'].lower().replace(".exe", "")
            
            # 使用精确匹配 (==) 而不是模糊匹配 (in)
            # 防止目标是 "code"，却意外关掉了名为 "vscode-server" 或其他包含该词的系统进程
            if target_name == current_proc_name:
                proc.terminate()  # 发送终止信号
                terminated_count += 1
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            # 忽略权限不足或已消失的进程
            continue

    # 5. 组装给大模型的标准化反馈
    if terminated_count > 0:
        return f"成功：已在 {system_type} 环境下关闭了 {terminated_count} 个 '{target_name}' 进程。"
    else:
        return f"提示：未找到正在运行的 '{target_name}' 进程，它可能已经被关闭。"




def execute_in_sandbox(bash_script: str, wsl_workspace_path: str = "/home/kl/AI_Workspace") -> str:
    """
    AI 助手的专属代码执行工具。
    将脚本写入 WSL 网络路径，并通过 wsl 命令拉取 Docker 容器执行。
    """
    # 1. 自动转换路径：将 WSL 路径 (/home/xxx/ai_workspace) 转为 Windows 网络路径
    # 请确保把你 WSL 的实际用户名替换到这里（如果你忘了，可以在 WSL 里敲 whoami 查看）
    wsl_distro_name = "Ubuntu" 
    win_unc_path = f"\\\\wsl.localhost\\{wsl_distro_name}{wsl_workspace_path.replace('/', '\\')}"
    
    # 2. 将 AI 生成的脚本保存到工作区
    script_name = "ai_agent_task.sh"
    script_path = os.path.join(win_unc_path, script_name)
    
    # 注意：必须使用 newline='\n' 保证 Linux 换行符
    try:
        with open(script_path, "w", encoding="utf-8", newline='\n') as f:
            f.write(bash_script)
    except Exception as e:
        return {"success": False, "error": f"无法写入文件到 {win_unc_path}: {e}"}

    # 3. 组装 Docker 命令（限制内存、CPU，挂载目录）
    # 组装 Docker 命令（去掉了 sudo，因为我们已经配置了免密）
    cmd = [
        "wsl", "-e", "docker", "run", "--rm",
        "-m", "2048m",                      
        "--cpus", "2.0",                   
        "-v", f"{wsl_workspace_path}:/workspace", 
        "-w", "/workspace",                
        "my-ai-sandbox:v1",                
        "bash", script_name                
    ]

    print(f"🚀 正在沙盒中执行任务...")

    try:
        # 【关键修复 1】明确指定 encoding='utf-8'，防止 Windows 默认用 GBK 解码报错
        result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8', timeout=300)
        
        if os.path.exists(script_path):
            os.remove(script_path)
            
        return f"任务执行{'成功' if result.returncode == 0 else '失败'};标准输出:\n{result.stdout.strip() if result.stdout else ''};标准错误:\n{result.stderr.strip() if result.stderr else ''}"
    except subprocess.TimeoutExpired as e:
        if os.path.exists(script_path):
            os.remove(script_path)
        
        # 因为我们指定了 encoding='utf-8'，这里的 e.stdout 直接就是字符串了
        partial_out = e.stdout if isinstance(e.stdout, str) else (e.stdout.decode('utf-8') if e.stdout else "")
        partial_err = e.stderr if isinstance(e.stderr, str) else (e.stderr.decode('utf-8') if e.stderr else "")
        
        return f"执行超时 (300秒)。\n超时前输出:\n{partial_out}\n错误日志:\n{partial_err}"
    



Config = {
    # OSS配置 (通过环境变量或config.json设置)
    'OSS_ACCESS_KEY_ID': os.getenv("ALI_OSS_ACCESS_KEY_ID", ""),
    'OSS_ACCESS_KEY_SECRET': os.getenv("ALI_OSS_ACCESS_KEY_SECRET", ""),
    'OSS_BUCKET_NAME': "ysnetdisk",
    'OSS_REGION': "cn-wulanchabu",
    'OSS_ENDPOINT': "https://oss-cn-wulanchabu.aliyuncs.com", # -internal
    
    'MAX_CONTENT_LENGTH': 1024 * 1024 * 1024,  # 最大上传文件大小 1GB
    
    # 预签名URL过期时间
    'UPLOAD_EXPIRE_TIME': timedelta(seconds=3600),  # 1小时
    'DOWNLOAD_EXPIRE_TIME': timedelta(seconds=1800),  # 30分钟
}

# 初始化OSS客户端
def get_oss_client():
    cfg = oss.config.load_default()
    cfg.credentials_provider = StaticCredentialsProvider(
        Config['OSS_ACCESS_KEY_ID'],
        Config['OSS_ACCESS_KEY_SECRET']
    )
    cfg.region = Config['OSS_REGION']
    cfg.endpoint = Config['OSS_ENDPOINT']
    return oss.Client(cfg)

def get_oss2_bucket():
    """用于生成下载URL的客户端（使用oss2库）"""
    import oss2
    auth = oss2.Auth(
        Config['OSS_ACCESS_KEY_ID'],
        Config['OSS_ACCESS_KEY_SECRET']
    )
    bucket = oss2.Bucket(
        auth,
        Config['OSS_ENDPOINT'],
        Config['OSS_BUCKET_NAME']
    )
    return bucket

def upload_file_and_get_url(local_file_path, user_id=None, filename=None, expires=3600):
    """
    上传本地图片到 OSS 的 tmp 目录，并返回预签名下载 URL。

    :param local_file_path: 本地图片文件路径
    :param user_id: 用户 ID（可选），若提供则路径为 users/{user_id}/tmp/，否则为公共 tmp/
    :param filename: 自定义文件名（可选），默认使用原文件名
    :param expires: 预签名 URL 有效期（秒），默认 3600 秒（1 小时）
    :return: 预签名下载 URL
    :raises FileNotFoundError: 本地文件不存在时抛出
    """
    # 检查本地文件是否存在
    if not os.path.isfile(local_file_path):
        raise FileNotFoundError(f"文件不存在: {local_file_path}")

    # 确定文件名
    if not filename:
        filename = os.path.basename(local_file_path)

    # 构造 OSS 对象键（key）
    if user_id:
        key = f"users/{user_id}/tmp/{filename}"
    else:
        key = f"tmp/{filename}"

    
    client = get_oss_client()

    # 上传文件（直接读取本地文件并上传）
    with open(local_file_path, 'rb') as f:
        client.put_object(oss.PutObjectRequest(
            bucket=Config['OSS_BUCKET_NAME'],
            key=key,
            body=f
        ))

    # 生成外网可访问的预签名下载 URL（使用 oss2 库）
    bucket = get_oss2_bucket()
    url = bucket.sign_url('GET', key, expires, slash_safe=True)
    # 处理 URL 中的特殊字符（参考原代码）
    url = url.replace('+', '%2B')

    return url



if __name__ == "__main__":
    # print(check_location())
    # text = read_pdf("D:/课件/机器学习课件/C05_贝叶斯分类与随机图模型.pdf")
    # text = text.replace("\n", "")
    # text = text.replace("西安交通大学人工智能学院魏平编写。课程资料，请勿外传", "")
    # text = text.replace("魏平", "")
    # text = text.replace("西安交通大学", "")
    # print(text)

    print(search_deep_memory("你好"))
